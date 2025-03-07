import streamlit as st
import numpy as np
import pandas as pd
from  sklearn.preprocessing import StandardScaler,LabelEncoder
from tensorflow.keras.models import load_model
import docx
from docx import Document
import pdfkit
import io
import base64
import os
import tempfile
import platform

# Custom styling
st.set_page_config(
    page_title="AI Disease Predictor",
    page_icon="🏥",
    layout="wide"
)

# Custom CSS with professional color scheme
st.markdown("""
    <style>
    .main-header {
        font-size: 3rem;
        color: rgb(236, 245, 254);
        text-align: center;
        padding: 1.5rem;
        font-weight: bold;
        text-shadow: 2px 2px 4px rgb(255, 255, 255);
        margin-bottom: 2rem;
        background: linear-gradient(to right, rgb(235, 244, 255), rgb(243, 248, 254));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        border-bottom: 2px solid rgb(6, 6, 6);
    }
    .sub-header {
        font-size: 1.5rem;
        color: #576574;
        text-align: center;
        margin-bottom: 2rem;
        font-weight: 500;
    }
    </style>
""", unsafe_allow_html=True)

# Updated headers with new styling
st.markdown('<h1 class="main-header">🏥 Disease Prediction System</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Professional Medical Diagnosis Assistant</p>', unsafe_allow_html=True)

# Add a separator
st.markdown("---")

# Show instructions popup on first load
if 'show_instructions' not in st.session_state:
    st.session_state.show_instructions = True

if st.session_state.show_instructions:
    with st.expander("📋 How to Use This Application", expanded=True):
        st.markdown("""
        ## Welcome to Disease Prediction System! 👋
        
        ### How to use this app:
        1. **Enter your symptoms** in the text box below (e.g., 'itching, skin rash')
        2. **Select matching symptoms** from the dropdown menu
        3. **View predictions** based on your selected symptoms
        4. **Get detailed information** about predicted diseases
        5. **Download results** as PDF or Word document if needed
        
        > 📝 **Note:** The more symptoms you enter, the more accurate the prediction will be. Please select at least 2 symptoms.
        
        **Disclaimer:** This app is for informational purposes only and should not replace professional medical advice.
        
        ❓ Click the X in the top right to close these instructions.
        """)
        if st.button("Got it!", key="dismiss_instructions"):
            st.session_state.show_instructions = False
            st.experimental_rerun()

# Separate data loading functions
def load_training_data():
    """Load only the data needed for model training"""
    try:
        data = pd.read_csv('dataset.csv')  # Switch to dataset.csv
        if 'Disease' not in data.columns:
            raise ValueError("Dataset is missing 'Disease' column.")
        if data['Disease'].isnull().all() or data['Disease'].empty:
            raise ValueError("Disease column contains no valid data.")
        return data
    except Exception as e:
        st.error(f"Error loading training data: {e}")
        return None

def load_disease_descriptions():
    """Load disease descriptions separately"""
    try:
        return pd.read_csv('symptom_Description.csv')
    except:
        st.error("Could not load descriptions file")
        return None

def load_disease_precautions():
    """Load disease precautions separately"""
    try:
        return pd.read_csv('symptom_precaution.csv')
    except:
        st.error("Could not load precautions file")
        return None

# Preprocess input features (updated for raw dataset with Symptom_1 to Symptom_17)
def preprocess_input_features(data):
    symptoms = [
        "itching", "skin_rash", "nodal_skin_eruptions", "dischromic_patches", "continuous_sneezing", "shivering",
        "chills", "watering_from_eyes", "stomach_pain", "acidity", "ulcers_on_tongue", "vomiting", "cough",
        "chest_pain", "yellowish_skin", "nausea", "loss_of_appetite", "abdominal_pain", "yellowing_of_eyes",
        "burning_micturition", "spotting_urination", "passage_of_gases", "internal_itching", "indigestion",
        "muscle_wasting", "patches_in_throat", "high_fever", "extra_marital_contacts", "fatigue", "weight_loss",
        "restlessness", "lethargy", "irregular_sugar_level", "blurred_and_distorted_vision", "obesity",
        "excessive_hunger", "increased_appetite", "polyuria", "sunken_eyes", "dehydration", "diarrhoea",
        "breathlessness", "family_history", "mucoid_sputum", "headache", "dizziness", "loss_of_balance",
        "lack_of_concentration", "stiff_neck", "depression", "irritability", "visual_disturbances", "back_pain",
        "weakness_in_limbs", "neck_pain", "weakness_of_one_body_side", "altered_sensorium", "dark_urine",
        "sweating", "muscle_pain", "mild_fever", "swelled_lymph_nodes", "malaise", "red_spots_over_body",
        "joint_pain", "pain_behind_the_eyes", "constipation", "toxic_look_(typhos)", "belly_pain", "yellow_urine",
        "receiving_blood_transfusion", "receiving_unsterile_injections", "coma", "stomach_bleeding",
        "acute_liver_failure", "swelling_of_stomach", "distention_of_abdomen", "history_of_alcohol_consumption",
        "fluid_overload", "phlegm", "blood_in_sputum", "throat_irritation", "redness_of_eyes", "sinus_pressure",
        "runny_nose", "congestion", "loss_of_smell", "fast_heart_rate", "rusty_sputum", "pain_during_bowel_movements",
        "pain_in_anal_region", "bloody_stool", "irritation_in_anus", "cramps", "bruising", "swollen_legs",
        "swollen_blood_vessels", "prominent_veins_on_calf", "weight_gain", "cold_hands_and_feets", "mood_swings",
        "puffy_face_and_eyes", "enlarged_thyroid", "brittle_nails", "swollen_extremeties", "abnormal_menstruation",
        "muscle_weakness", "anxiety", "slurred_speech", "palpitations", "drying_and_tingling_lips", "knee_pain",
        "hip_joint_pain", "swelling_joints", "painful_walking", "movement_stiffness", "spinning_movements",
        "unsteadiness", "pus_filled_pimples", "blackheads", "scurring"
    ]  # 121 symptoms

    x = np.zeros((len(data), len(symptoms)))
    
    # Handle raw dataset with Symptom_1 to Symptom_17 columns
    for idx, row in data.iterrows():
        for col in [f'Symptom_{i}' for i in range(1, 18)]:  # Symptom_1 to Symptom_17
            symptom = row[col]
            if pd.notna(symptom):  # Check for non-null values
                symptom = symptom.strip()  # Strip spaces from string symptoms
                if symptom in symptoms:
                    x[idx, symptoms.index(symptom)] = 1

    return x

# Load model and fit scaler on training data
@st.cache_resource
def load_model_and_scaler():
    try:
        model = load_model('model2.h5')
        training_data = load_training_data()
        if training_data is None:
            return None, None, None
        x = preprocess_input_features(training_data)
        scaler = StandardScaler()
        scaler.fit(x)  # Fit scaler once on training data
        return model, scaler, x  # Return x for label encoder
    except Exception as e:
        st.error(f"Error loading model: {e}")
        return None, None, None

# Prediction function
def predict_disease(user_symptoms, model, scaler, symptoms, label_encoder):
    input_data = np.zeros((1, len(symptoms)))
    
    # Convert symptoms to binary vector
    for symptom in user_symptoms:
        if symptom in symptoms:
            input_data[0, symptoms.index(symptom)] = 1
    
    # Debug: Print user input symptoms
    st.write("🔍 **Debug - User Selected Symptoms:**", user_symptoms)
    st.write("🛠 **Binary Encoded Input Data:**", input_data)
    st.write(f"Input data shape: {input_data.shape}")  # Debug shape

    # Standardize input
    standardized_input = scaler.transform(input_data)

    # Debug: Print transformed inputs after scaling
    st.write("📊 **Standardized Input Data:**", standardized_input)

    predictions = model.predict(standardized_input)[0]

    # Debug: Print raw model outputs
    st.write("📈 **Raw Model Predictions (Probabilities):**", predictions)

    # Adjust confidence threshold
    CONFIDENCE_THRESHOLD = 0.1  # Increased threshold to avoid always same prediction
    
    # Get all predictions with their indices
    disease_predictions = [(i, prob) for i, prob in enumerate(predictions)]
    
    # Sort by probability and get top 5
    disease_predictions.sort(key=lambda x: x[1], reverse=True)
    top_predictions = disease_predictions[:5]
    
    # Filter predictions above threshold
    valid_predictions = [(idx, prob) for idx, prob in top_predictions if prob >= CONFIDENCE_THRESHOLD]
    
    # If no predictions meet threshold, return top 3 anyway
    if not valid_predictions:
        valid_predictions = top_predictions[:3]

    # Convert to disease names with probabilities
    results = []
    for idx, prob in valid_predictions:
        disease_name = label_encoder.inverse_transform([idx])[0]
        results.append((disease_name, prob))

    # Debug: Print the final predicted diseases
    st.write("✅ **Final Predicted Diseases & Confidence Scores:**", results)

    return results

st.cache_resource.clear()  
model, scaler, x = load_model_and_scaler()
if model is None or scaler is None:
    st.stop()  # Stop execution if there's an error loading the model

data = load_training_data()
if data is None:
    st.stop()  
    
label_encoder = LabelEncoder()
try:
    label_encoder.fit(data['Disease'])
except Exception as e:
    st.error(f"Error fitting label encoder: {e}")
    st.stop()

# Prepare symptom list for autocomplete (121 symptoms)
symptoms_list = [
    "itching", "skin_rash", "nodal_skin_eruptions", "dischromic_patches", "continuous_sneezing", "shivering",
    "chills", "watering_from_eyes", "stomach_pain", "acidity", "ulcers_on_tongue", "vomiting", "cough",
    "chest_pain", "yellowish_skin", "nausea", "loss_of_appetite", "abdominal_pain", "yellowing_of_eyes",
    "burning_micturition", "spotting_urination", "passage_of_gases", "internal_itching", "indigestion",
    "muscle_wasting", "patches_in_throat", "high_fever", "extra_marital_contacts", "fatigue", "weight_loss",
    "restlessness", "lethargy", "irregular_sugar_level", "blurred_and_distorted_vision", "obesity",
    "excessive_hunger", "increased_appetite", "polyuria", "sunken_eyes", "dehydration", "diarrhoea",
    "breathlessness", "family_history", "mucoid_sputum", "headache", "dizziness", "loss_of_balance",
    "lack_of_concentration", "stiff_neck", "depression", "irritability", "visual_disturbances", "back_pain",
    "weakness_in_limbs", "neck_pain", "weakness_of_one_body_side", "altered_sensorium", "dark_urine",
    "sweating", "muscle_pain", "mild_fever", "swelled_lymph_nodes", "malaise", "red_spots_over_body",
    "joint_pain", "pain_behind_the_eyes", "constipation", "toxic_look_(typhos)", "belly_pain", "yellow_urine",
    "receiving_blood_transfusion", "receiving_unsterile_injections", "coma", "stomach_bleeding",
    "acute_liver_failure", "swelling_of_stomach", "distention_of_abdomen", "history_of_alcohol_consumption",
    "fluid_overload", "phlegm", "blood_in_sputum", "throat_irritation", "redness_of_eyes", "sinus_pressure",
    "runny_nose", "congestion", "loss_of_smell", "fast_heart_rate", "rusty_sputum", "pain_during_bowel_movements",
    "pain_in_anal_region", "bloody_stool", "irritation_in_anus", "cramps", "bruising", "swollen_legs",
    "swollen_blood_vessels", "prominent_veins_on_calf", "weight_gain", "cold_hands_and_feets", "mood_swings",
    "puffy_face_and_eyes", "enlarged_thyroid", "brittle_nails", "swollen_extremeties", "abnormal_menstruation",
    "muscle_weakness", "anxiety", "slurred_speech", "palpitations", "drying_and_tingling_lips", "knee_pain",
    "hip_joint_pain", "swelling_joints", "painful_walking", "movement_stiffness", "spinning_movements",
    "unsteadiness", "pus_filled_pimples", "blackheads", "scurring"
]  # 121 symptoms

# User input with auto-completion feature
typed_symptoms = st.text_input("Enter your symptoms (e.g., 'itching, skin rash'):")
if typed_symptoms:
    symptom_keywords = [s.strip().lower() for s in typed_symptoms.split(",")]
    
    # Find matching symptoms
    matched_symptoms = [symptom for symptom in symptoms_list if any(keyword in symptom.lower() for keyword in symptom_keywords)]
    
    if matched_symptoms:
        selected_symptoms = st.multiselect("Select your symptoms:", options=matched_symptoms)
        
        if selected_symptoms:
            if len(selected_symptoms) < 2:
                st.warning("Please select at least 2 symptoms for more accurate prediction.")
            else:
                predicted_diseases = predict_disease(selected_symptoms, model, scaler, symptoms_list, label_encoder)
                
                if predicted_diseases:
                    st.subheader("🔍 Predicted Diseases:")
                    
                    # Load additional information only when needed
                    descriptions_data = load_disease_descriptions()
                    precautions_data = load_disease_precautions()
                    
                    # Store prediction results for export
                    prediction_results = []
                    
                    for disease, probability in predicted_diseases:
                        confidence = probability * 100
                        
                        # Assign emoji based on confidence
                        if confidence >= 70:
                            confidence_color = "green"
                            emoji = "✅ "
                        elif confidence >= 40:
                            confidence_color = "orange"
                            emoji = "⚠️ "
                        else:
                            confidence_color = "red"
                            emoji = "❓ "
                        
                        result_text = f"{emoji}**{disease}** (::{confidence_color}[Confidence: {confidence:.1f}%]::)"
                        st.markdown(result_text)
                        
                        # Store for export
                        prediction_results.append({
                            "disease": disease,
                            "confidence": confidence,
                            "emoji": emoji
                        })
                        
                        with st.expander(f"More information about {disease}"):
                            # Display description if available
                            description_text = ""
                            precautions_text = []
                            
                            if descriptions_data is not None:
                                disease_desc = descriptions_data[
                                    descriptions_data['Disease'] == disease
                                ]
                                if not disease_desc.empty:
                                    st.markdown("### Description")
                                    desc_text = disease_desc.iloc[0]['Description']
                                    st.write(desc_text)
                                    description_text = desc_text
                            
                            # Display precautions if available
                            if precautions_data is not None:
                                disease_prec = precautions_data[
                                    precautions_data['Disease'] == disease
                                ]
                                if not disease_prec.empty:
                                    st.markdown("### Recommended Precautions")
                                    prec = disease_prec.iloc[0]
                                    for i in range(1, 5):
                                        precaution = prec.get(f'Precaution_{i}')
                                        if pd.notna(precaution) and precaution != '':
                                            st.write(f"{i}. {precaution.capitalize()}")
                                            precautions_text.append(precaution.capitalize())
                            
                            # Add to stored results
                            prediction_results[-1]["description"] = description_text
                            prediction_results[-1]["precautions"] = precautions_text
                    
                    # Add download options
                    st.markdown("---")
                    st.subheader("📥 Download Prediction Results")
                    
                    # Create columns for the buttons
                    col1, col2 = st.columns(2)
                    
                    # Function to generate DOCX
                    def generate_docx():
                        doc = Document()
                        doc.add_heading('Disease Prediction Results', 0)
                        
                        # Add patient symptoms
                        doc.add_heading('Symptoms Reported', level=1)
                        for symptom in selected_symptoms:
                            doc.add_paragraph(f"• {symptom}", style='List Bullet')
                        
                        # Add predicted diseases
                        doc.add_heading('Predicted Diseases', level=1)
                        for result in prediction_results:
                            emoji_map = {"✅ ": "[HIGH] ", "⚠️ ": "[MEDIUM] ", "❓ ": "[LOW] "}
                            emoji_text = emoji_map.get(result["emoji"], "")
                            p = doc.add_paragraph(f"{emoji_text}{result['disease']} - Confidence: {result['confidence']:.1f}%")
                            p.style = 'Heading 2'
                            
                            # Add description
                            if result["description"]:
                                doc.add_heading('Description', level=3)
                                doc.add_paragraph(result["description"])
                            
                            # Add precautions
                            if result["precautions"]:
                                doc.add_heading('Recommended Precautions', level=3)
                                for i, precaution in enumerate(result["precautions"]):
                                    doc.add_paragraph(f"{i+1}. {precaution}")
                            
                            # Add separator between diseases
                            doc.add_paragraph("---")
                        
                        # Add disclaimer
                        doc.add_heading('Disclaimer', level=1)
                        doc.add_paragraph("This prediction is generated by an AI model and should not replace professional medical advice. Please consult a healthcare professional for proper diagnosis and treatment.")
                        
                        return doc
                    
                    # Function to check wkhtmltopdf installation with improved error handling
                    def check_wkhtmltopdf_installation():
                        """Check if wkhtmltopdf is installed and provide installation instructions if not"""
                        try:
                            import subprocess
                            
                            # Handle different platforms
                            if platform.system() == "Windows":
                                # Try to find wkhtmltopdf in common locations
                                common_paths = [
                                    r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe",
                                    r"C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe"
                                ]
                                
                                for path in common_paths:
                                    if os.path.exists(path):
                                        # Configure pdfkit with the path
                                        config = pdfkit.configuration(wkhtmltopdf=path)
                                        return config
                                        
                                # Try using PATH
                                try:
                                    result = subprocess.run(['wkhtmltopdf', '-V'], 
                                                           stdout=subprocess.PIPE, 
                                                           stderr=subprocess.PIPE,
                                                           check=False)
                                    if result.returncode == 0:
                                        return None  # Use default configuration
                                except:
                                    pass
                                    
                            elif platform.system() == "Linux":
                                try:
                                    result = subprocess.run(['which', 'wkhtmltopdf'], 
                                                           stdout=subprocess.PIPE, 
                                                           stderr=subprocess.PIPE,
                                                           check=False)
                                    if result.returncode == 0:
                                        return None  # Use default configuration
                                except:
                                    pass
                                    
                            elif platform.system() == "Darwin":  # macOS
                                try:
                                    result = subprocess.run(['which', 'wkhtmltopdf'], 
                                                           stdout=subprocess.PIPE, 
                                                           stderr=subprocess.PIPE,
                                                           check=False)
                                    if result.returncode == 0:
                                        return None  # Use default configuration
                                except:
                                    pass
                                    
                            # If we get here, wkhtmltopdf was not found
                            return False
                            
                        except Exception as e:
                            st.error(f"Error checking wkhtmltopdf: {e}")
                            return False

                    # Function to generate PDF with improved error handling
                    def generate_pdf():
                        # First create HTML content
                        html_content = f"""
                        <html>
                        <head>
                            <meta charset="UTF-8">
                            <style>
                                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                                h1 {{ color: #2c3e50; border-bottom: 1px solid #eee; padding-bottom: 10px; }}
                                h2 {{ color: #3498db; margin-top: 20px; }}
                                h3 {{ color: #7f8c8d; }}
                                .symptom {{ padding: 5px; }}
                                .disease {{ margin-bottom: 15px; padding: 10px; border-left: 4px solid #3498db; background-color: #f9f9f9; }}
                                .high {{ border-left-color: #2ecc71; }}
                                .medium {{ border-left-color: #f39c12; }}
                                .low {{ border-left-color: #e74c3c; }}
                                .disclaimer {{ font-style: italic; margin-top: 30px; color: #7f8c8d; }}
                            </style>
                        </head>
                        <body>
                            <h1>Disease Prediction Results</h1>
                            <h2>Symptoms Reported</h2>
                        """
                        
                        # Add symptoms
                        for symptom in selected_symptoms:
                            html_content += f"<div class='symptom'>• {symptom}</div>"
                        
                        html_content += "<h2>Predicted Diseases</h2>"
                        
                        # Add diseases
                        for result in prediction_results:
                            confidence_class = "high" if result["confidence"] >= 70 else ("medium" if result["confidence"] >= 40 else "low")
                            emoji_text = "HIGH CONFIDENCE: " if result["confidence"] >= 70 else ("MEDIUM CONFIDENCE: " if result["confidence"] >= 40 else "LOW CONFIDENCE: ")
                            
                            html_content += f"<div class='disease {confidence_class}'>"
                            html_content += f"<h3>{emoji_text}{result['disease']} - {result['confidence']:.1f}%</h3>"
                            
                            if result["description"]:
                                html_content += "<h4>Description:</h4>"
                                html_content += f"<p>{result['description']}</p>"
                            
                            if result["precautions"]:
                                html_content += "<h4>Recommended Precautions:</h4>"
                                html_content += "<ol>"
                                for precaution in result["precautions"]:
                                    html_content += f"<li>{precaution}</li>"
                                html_content += "</ol>"
                            
                            html_content += "</div>"
                        
                        # Add disclaimer
                        html_content += """
                            <div class="disclaimer">
                                <h2>Disclaimer</h2>
                                <p>This prediction is generated by an AI model and should not replace professional medical advice. 
                                Please consult a healthcare professional for proper diagnosis and treatment.</p>
                            </div>
                        </body>
                        </html>
                        """
                        
                        # Create a temporary HTML file with UTF-8 encoding
                        with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as f:
                            f.write(html_content)
                            temp_html = f.name
                        
                        try:
                            # Get wkhtmltopdf configuration
                            config = check_wkhtmltopdf_installation()
                            
                            # If wkhtmltopdf is not available
                            if config is False:
                                st.error("wkhtmltopdf is not installed or not in your PATH")
                                st.info("""
                                ### Please install wkhtmltopdf:
                                
                                1. Download from [wkhtmltopdf.org](https://wkhtmltopdf.org/downloads.html)
                                2. Install the package
                                3. Make sure it's in your system PATH
                                4. Restart this application
                                """)
                                # Clean up the temp file
                                os.unlink(temp_html)
                                return None
                                
                            # Convert HTML to PDF with options
                            options = {
                                'encoding': "UTF-8",
                                'page-size': 'A4',
                                'margin-top': '0.75in',
                                'margin-right': '0.75in',
                                'margin-bottom': '0.75in',
                                'margin-left': '0.75in',
                                'quiet': '',
                            }
                            
                            if config:  # If we have a specific configuration
                                pdf_data = pdfkit.from_file(temp_html, False, options=options, configuration=config)
                            else:  # Use default configuration
                                pdf_data = pdfkit.from_file(temp_html, False, options=options)
                                
                            # Clean up the temp file
                            os.unlink(temp_html)
                            return pdf_data
                            
                        except Exception as e:
                            st.error(f"Error generating PDF: {str(e)}")
                            # Clean up on error
                            if os.exists(temp_html):
                                os.unlink(temp_html)
                            return None

                    # Generate and offer DOCX for download
                    with col1:
                        if st.button("Generate Word Document"):
                            with st.spinner("Generating Word document..."):
                                try:
                                    doc = generate_docx()
                                    bio = io.BytesIO()
                                    doc.save(bio)
                                    bio.seek(0)
                                    
                                    st.success("Word document generated successfully!")
                                    st.download_button(
                                        label="Download Word Document",
                                        data=bio,
                                        file_name="disease_prediction.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="word_download"
                                    )
                                except Exception as e:
                                    st.error(f"Error generating Word document: {str(e)}")
                                    st.info("Please try again or contact support if the issue persists.")

                    # Modified PDF generation button code
                    with col2:
                        if st.button("Generate PDF"):
                            with st.spinner("Generating PDF..."):
                                try:
                                    pdf_data = generate_pdf()
                                    if pdf_data:
                                        st.success("PDF generated successfully!")
                                        st.download_button(
                                            label="Download PDF",
                                            data=pdf_data,
                                            file_name="disease_prediction.pdf",
                                            mime="application/pdf",
                                            key="pdf_download"
                                        )
                                    else:
                                        # Offer Word document as fallback
                                        st.warning("PDF generation failed. Would you like to generate a Word document instead?")
                                        if st.button("Generate Word Document Instead", key="fallback_word"):
                                            with st.spinner("Generating Word document..."):
                                                try:
                                                    doc = generate_docx()
                                                    bio = io.BytesIO()
                                                    doc.save(bio)
                                                    bio.seek(0)
                                                    
                                                    st.success("Word document generated successfully!")
                                                    st.download_button(
                                                        label="Download Word Document",
                                                        data=bio,
                                                        file_name="disease_prediction.docx",
                                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                        key="fallback_word_download"
                                                    )
                                                except Exception as e:
                                                    st.error(f"Error generating Word document: {str(e)}")
                                except Exception as e:
                                    st.error(f"PDF generation error: {str(e)}")
                                    st.info("If you've installed wkhtmltopdf but still see this error, please check that it's properly installed and in your system PATH.")
                else:
                    st.warning("No confident predictions found. Please check your symptoms or add more symptoms.")
        else:
            st.write("Please select at least one symptom.")
    else:
        st.write("No matching symptoms found.")
else:
    st.write("Start typing symptoms to get suggestions.")

